from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE

OUT = "output/Decoupled_Vector_Runahead_论文解读与复现指南.docx"
BLUE = "245A7A"; DARK = "17324D"; LIGHT = "EAF1F5"; PALE = "F5F8FA"; GOLD = "B07A24"; RED = "9B2C2C"; GRAY = "66737F"

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.top_margin = sec.bottom_margin = Inches(0.78)
sec.left_margin = sec.right_margin = Inches(0.82)
sec.header_distance = Inches(0.35); sec.footer_distance = Inches(0.35)

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = tcPr.find(qn('w:shd'))
    if shd is None: shd = OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'), fill)

def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None: tcMar = OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for m,v in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node = tcMar.find(qn('w:'+m))
        if node is None: node=OxmlElement('w:'+m); tcMar.append(node)
        node.set(qn('w:w'), str(v)); node.set(qn('w:type'),'dxa')

def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr(); x=OxmlElement('w:tblHeader'); x.set(qn('w:val'),'true'); trPr.append(x)

def set_repeat_heading(p):
    pPr=p._p.get_or_add_pPr(); x=OxmlElement('w:keepNext'); pPr.append(x)

def font(run, size=10.5, bold=False, color=DARK, italic=False, name='STSong'):
    run.font.name=name; run._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'),name)
    run.font.size=Pt(size); run.bold=bold; run.italic=italic; run.font.color.rgb=RGBColor.from_string(color)
    return run

styles=doc.styles
normal=styles['Normal']; normal.font.name='STSong'; normal._element.rPr.rFonts.set(qn('w:eastAsia'),'STSong'); normal.font.size=Pt(10.5); normal.font.color.rgb=RGBColor.from_string(DARK)
normal.paragraph_format.space_after=Pt(5); normal.paragraph_format.line_spacing=1.23
for name,size,color,before,after in [('Title',27,DARK,0,8),('Subtitle',13,GRAY,0,8),('Heading 1',17,BLUE,15,7),('Heading 2',13.5,DARK,11,5),('Heading 3',11.5,BLUE,8,3)]:
    s=styles[name]; s.font.name='STSong'; s._element.rPr.rFonts.set(qn('w:eastAsia'),'STSong'); s.font.size=Pt(size); s.font.color.rgb=RGBColor.from_string(color); s.font.bold=(name!='Subtitle')
    s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True

for sname, fill, border in [('DVR Callout', LIGHT, BLUE), ('DVR Warning','FCEEEE',RED), ('DVR Note','F7F1E5',GOLD)]:
    if sname not in styles:
        s=styles.add_style(sname,WD_STYLE_TYPE.PARAGRAPH); s.base_style=normal
    s=styles[sname]; s.paragraph_format.left_indent=Inches(.18); s.paragraph_format.right_indent=Inches(.18); s.paragraph_format.space_before=Pt(6); s.paragraph_format.space_after=Pt(7)

def border_paragraph(p, color=BLUE, fill=None):
    pPr=p._p.get_or_add_pPr(); pbdr=OxmlElement('w:pBdr'); left=OxmlElement('w:left'); left.set(qn('w:val'),'single'); left.set(qn('w:sz'),'18'); left.set(qn('w:space'),'8'); left.set(qn('w:color'),color); pbdr.append(left); pPr.append(pbdr)
    if fill:
        shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); pPr.append(shd)

def callout(label,text,kind='callout'):
    style={'callout':'DVR Callout','warning':'DVR Warning','note':'DVR Note'}[kind]; color={'callout':BLUE,'warning':RED,'note':GOLD}[kind]; fill={'callout':LIGHT,'warning':'FCEEEE','note':'F7F1E5'}[kind]
    p=doc.add_paragraph(style=style); border_paragraph(p,color,fill); font(p.add_run(label+'｜'),10.5,True,color); font(p.add_run(text),10.5,False,DARK); return p

def heading(text,level=1):
    p=doc.add_paragraph(text,style=f'Heading {level}'); set_repeat_heading(p); return p

def para(text,boldlead=None):
    p=doc.add_paragraph()
    if boldlead and text.startswith(boldlead): font(p.add_run(boldlead),10.5,True,BLUE); text=text[len(boldlead):]
    font(p.add_run(text)); return p

def bullet(text, level=0):
    p=doc.add_paragraph(style='List Bullet' if level==0 else 'List Bullet 2'); p.paragraph_format.left_indent=Inches(.3+.2*level); p.paragraph_format.first_line_indent=Inches(-.16); p.paragraph_format.space_after=Pt(3); font(p.add_run(text)); return p

def number(text):
    p=doc.add_paragraph(style='List Number'); p.paragraph_format.left_indent=Inches(.34); p.paragraph_format.first_line_indent=Inches(-.18); p.paragraph_format.space_after=Pt(4); font(p.add_run(text)); return p

def table(headers, rows, widths=None, small=9.2):
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False; t.style='Table Grid'
    if widths is None: widths=[6.85/len(headers)]*len(headers)
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.width=Inches(widths[i]); set_cell_shading(c,BLUE); set_cell_margins(c); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(p.add_run(h),small,True,'FFFFFF')
    set_repeat_table_header(t.rows[0])
    for ri,row in enumerate(rows):
        cells=t.add_row().cells
        for i,val in enumerate(row):
            c=cells[i]; c.width=Inches(widths[i]); set_cell_margins(c); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if ri%2: set_cell_shading(c,PALE)
            p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0); font(p.add_run(str(val)),small,False,DARK)
    doc.add_paragraph().paragraph_format.space_after=Pt(1)
    return t

def pagebreak(): doc.add_page_break()

# Header/footer
hp=sec.header.paragraphs[0]; hp.alignment=WD_ALIGN_PARAGRAPH.RIGHT; font(hp.add_run('DVR｜论文解读与复现指南'),8.5,True,GRAY)
fp=sec.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
font(fp.add_run('Decoupled Vector Runahead · 复现工作文档  '),8,False,GRAY)
fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); fp._p.append(fld)

# Cover
p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(85); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(p.add_run('MICRO 2023 · 论文复现手册'),10,True,GOLD)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(p.add_run('Decoupled Vector Runahead'),29,True,DARK)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(p.add_run('解耦向量超前执行：从论文机制到 Sniper 实现'),15,False,BLUE)
p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(24); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(p.add_run('中文解读 · 微体系结构拆解 · 分阶段复现路线 · 验收清单'),11,False,GRAY)
p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(78); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(p.add_run('面向：计算机体系结构研究 / 硬件预取 / 不规则访存'),10,False,GRAY)
callout('一句话结论','DVR 把“等 ROB 堵满才超前执行”改成“检测到可向量化的间接访存链就启动轻量子线程”，用最多 128 个标量等价 lane 提前完成依赖链并把数据带入 L1。')
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(36); font(p.add_run('依据论文：Naithani et al., MICRO ’23, pp. 17–31'),9,False,GRAY)
pagebreak()

heading('使用说明与复现结论',1)
callout('先说结论','这不是一个“加一个普通预取器”即可复现的工作。严格复现需要改动 Sniper 的流水线、寄存器映射、执行端口仲裁、LSQ/MSHR、缓存统计与分支控制；建议先完成事件驱动的功能模型，再逐步加入周期级资源竞争。')
para('本文把论文中的“明确描述”与“复现时必须自行决策的缺口”分开。标注为【论文给定】的参数可直接对齐；标注为【工程假设】的部分必须写入实验日志，并做灵敏度分析。')
heading('最小成功标准',2)
bullet('功能：能识别 stride load → 传播 taint → 找到最后一个依赖 load（FLR）→ 生成不越过预测循环边界的预取。')
bullet('机制：DVR 子线程与主线程并行，只有主线程无就绪指令占用某执行端口时，DVR 才使用该端口。')
bullet('效果：在至少一个两级间接访问 microbenchmark 上，DVR 提高平均并发 miss 数并降低主线程 demand-load 延迟。')
bullet('论文级：对齐 13 个 benchmark、5 个 GAP 图输入、500M ROI 指令、基线参数和图 7–12 的趋势。')
heading('可信度边界',2)
callout('重要','截至本指南编写时，从论文、作者页面和公开检索结果中未找到与 MICRO ’23 DVR 对应的官方 artifact/代码仓库。因此，位级、周期级完全复现存在不可消除的实现歧义；最现实的目标是“机制复现 + 趋势复现”，再联系作者索取 patch、benchmark 版本和运行脚本。','warning')
table(['复现层级','目标','预计难度','推荐顺序'],[
('L0 论文核对','复算参数、结构开销、图表趋势','低','第 1 步'),('L1 功能原型','正确生成预取地址；不建模全部争用','中','第 2 步'),('L2 周期模型','端口、寄存器、MSHR、缓存与主线程竞争','高','第 3 步'),('L3 论文级实验','完整 workload/input 与消融、敏感性','很高','第 4 步')],[1.35,2.85,1.05,1.6])

heading('1. 论文要解决什么问题',1)
para('不规则程序常见访问链 A[i] → B[f(A[i])] → C[g(B[…])]。第一步 A[i] 有规律，但 B、C 的地址只有前一级 load 返回后才能算出。传统 stride prefetcher 只能覆盖 A；经典 runahead 虽会预执行未来指令，却受制于 ROB 满触发、前端带宽和单条依赖链的长延迟。')
heading('1.1 Vector Runahead 为什么仍然不够',2)
bullet('VR 把多个未来循环迭代的同一条指令重排并向量化，用 gather 同时发出多条独立 miss。')
bullet('但 VR 仍以 full-ROB stall 为触发条件。ROB 越大、分支误预测越频繁，越难触发。论文图 2 显示 ROB 从 128 增至 512 时，full-ROB stall 时间从约 51% 降到约 5%。')
bullet('VR 不推断真实 loop bound，短内层循环可能越界预取，污染 cache、浪费带宽，甚至降速。')
heading('1.2 DVR 的四个关键创新',2)
table(['创新','直觉','解决的问题'],[
('解耦子线程','检测到合适 stride 后主动启动，不等 ROB 满','触发太晚/太少'),('Discovery Mode','选最内层 stride，确认依赖 load，推断剩余迭代数','错误触发与越界预取'),('SIMT 式向量化','16×AVX-512，每向量 8 个 64-bit lane，共 128 lane','依赖链下仍暴露巨大 MLP'),('Nested DVR','聚合多个内循环 invocation 的迭代','短内循环凑不满 lane'),('分歧/汇合','按目标 PC 分组 mask，用栈依次执行路径','链内存在条件分支')],[1.25,2.5,3.1])
callout('核心洞察','DVR 不试图让 helper thread 算得“正确”，只要它对未来 load 地址足够准确、足够早，并且不修改体系结构状态即可。')

heading('2. 从触发到终止：完整执行流程',1)
number('Stride detector（32 项 RPT）在主线程中识别稳定的规则 load 及 stride。')
number('进入 Discovery Mode，跟随主线程执行一轮；如果发现更内层的 stride load，就切换候选并重置 VTT/FLR。')
number('VTT 从候选 stride load 的目的寄存器开始传播 taint；任何地址输入被 taint 的 load 都会更新 FLR。')
number('Loop-bound detector 追踪 backward branch 及其 compare；比较 Discovery 前后寄存器映射，推断 bound、increment 和剩余迭代数。')
number('主线程再次到达候选 stride PC 时，初始化 VRAT 并启动 in-order DVR 子线程。')
number('Vectorizer 把候选 load 和所有 tainted dependents 扩成最多 16 条 AVX-512 指令，即 128 个标量等价 lane；尾部 lane 用 mask 关闭。')
number('VIR 按序发射；只有相同执行端口上主线程没有 ready 指令时，DVR 才可用端口。vector gather 在 LSQ 中拆为标量 load，各自占用 MSHR。')
number('若 branch lane 产生不同 next PC，则按目标 PC 分组，把 PC+128-bit mask 压入 8 项 reconvergence stack。')
number('到达 FLR，或分歧情况下到下一次 stride PC，或执行 200 条指令超时，DVR 终止并释放临时资源。')
number('主线程稍后访问同一 cache line，理想情况下成为 L1 hit；之后再次遇到 stride 可重新发现和启动。')

heading('3. Discovery Mode：真正决定准确率的部分',1)
heading('3.1 最内层 stride 选择',2)
para('RPT 为每个 stride-load PC 记录地址历史、stride 和置信度。Discovery 期间用“一位/项”的 seen 向量记录出现过的 stride load；若某个新的 stride PC 在当前候选再次出现前已经出现两次，说明它处于更内层循环，于是改用它作为候选。')
heading('3.2 VTT 与 FLR',2)
table(['事件','VTT 更新','FLR 更新'],[
('进入 Discovery','仅候选 stride load 的 dst=1','清零'),('普通指令 src 有 taint','dst=1','不变'),('dst 原为 taint，但所有 src 不 taint','dst=0','不变'),('load 的地址输入有 taint','按数据流传播','写入该 load PC；后出现者覆盖')],[2.0,2.5,2.35])
callout('实现提示','FLR 只是“Discovery 这一轮中最后看到的 tainted-address load”，不是静态 CFG 意义上的唯一末端。分支复杂时，论文允许执行到下一次 stride PC，以探索分歧路径。','note')
heading('3.3 Loop bound 推断',2)
para('每次 FLR 更新时清零 LCR 和 SBB。SBB=0 时，首个 compare 把源/目的架构寄存器 ID 写入 LCR；随后若 branch 使用该 compare 结果，且 taken target ≤ stride PC，则锁定该 compare。Discovery 进入/退出各保存一次架构寄存器映射快照：一个 compare 输入保持不变则视为 bound，另一个发生变化，其差值视为 increment。若匹配失败，论文默认最多运行 128 lane。')
table(['边界情况','建议实现'],[
('递减循环','允许 increment<0；剩余数用方向感知的整除/向上取整'),('非单位步长','用观测差值，不假设 +1'),('多出口/break','保留 200 指令 timeout；记录 timeout 次数'),('bound 推断失败','回退 128，但单独统计 overfetch'),('剩余数 < 128','仅激活有效 lane，其余 mask=0')],[1.75,5.1])

heading('4. 子线程数据通路与状态',1)
heading('4.1 VRAT：子线程寄存器映射',2)
para('DVR 与主线程共享标量/向量物理寄存器文件，但使用独立 VRAT。启动时为主线程可见的架构寄存器复制/分配新标量物理寄存器，使 helper 与主线程解耦。某架构寄存器可在所有 lane 指向同一标量物理寄存器，也可为 16 个向量组分别指向不同向量物理寄存器。')
bullet('首次遇到 vectorized source 而 dst 尚未 vectorized：为 dst 分配 16 个向量物理寄存器。')
bullet('vectorized dst 将被纯标量指令覆盖：切换成标量物理寄存器。')
bullet('覆盖后不再作为当前 VIR source：立即释放；若仍是 source，用 dead-source 位延迟到 execute 后释放。')
heading('4.2 VIR：极简 in-order 发射',2)
table(['字段/行为','论文配置','复现关注点'],[
('活跃 mask','128 bit','fault、FP、分歧、尾部 lane 均可关闭'),('issued / executed','各 16 bit','对应 16 个 AVX-512 副本'),('uop + imm','64 bit','需映射到 Sniper uop 表示'),('dst/src1/src2','9×16 / 10×16 / 10×16 bit','论文的紧凑编码与模拟器对象不必等大小'),('端口优先级','主线程优先','逐端口、逐周期仲裁'),('gather','LSQ 中拆标量 load','每条 miss 独立请求并受 MSHR 限制')],[1.75,1.55,3.55])
heading('4.3 分支分歧与汇合',2)
para('对所有活跃 lane 的 branch 结果求 next PC 分组。若不一致，为每个目标生成 mask+PC，压栈后先执行一组；论文把汇合点简化为 DVR 终止点。若分歧恰好落在 8-lane AVX-512 组边界，可继续用不同标量映射；若组内 lane 使用不同标量，必须把该目的映射升级为向量寄存器并复制标量值。')
callout('易错点','不要用主线程 branch predictor 的错误路径状态直接污染 DVR 的 lane mask。DVR 的控制流是“按每个 lane 的真实/模拟结果分组”，而不是单一 PC 的普通预测执行。','warning')

heading('5. Nested Vector Runahead',1)
para('当内循环剩余迭代少于 64 时，单次内循环无法提供足够 MLP。NDM 跨多个外循环迭代收集内循环入口和各自 bound，最终把总计最多 128 个内循环元素铺进 DVR lane。')
number('保存原内循环的 LCR 源寄存器、increment（IR）与 inner stride-load PC（ILR）。')
number('反转内循环 backward branch 方向，跳出其后续迭代；NDM 子线程先标量执行外层代码。')
number('若 200 条指令内发现地址早于 ILR 的 outer stride load，将其向量化 16 倍并传播 outer taint。')
number('抵达每个 inner stride load 时，从向量化 LCR 输入和 IR 计算该 invocation 的迭代数。')
number('按外层 lane 顺序收集 inner 起始地址，累计至 128；超出部分丢弃。')
number('重新装填内层 DVR 的 VRAT/向量寄存器，从 inner stride load 到 FLR 执行常规 DVR。')
callout('建议','先不要实现 NDM。先做单层 DVR 并证明准确率/MLP，再加入 NDM 做消融。否则 bug 很难区分来自 bound、外层 taint、lane flatten 还是寄存器回溯。','note')

heading('6. 论文给定的硬件参数',1)
table(['结构','配置/位数','论文开销'],[
('Stride detector','32 项；PC 48 + prev addr 48 + stride 16 + counter 2 + inner 1 bit/项','460 B'),('VRAT','16 项；每项 16×9-bit register id','288 B'),('VIR','mask 128；issued 16；executed 16；uop/imm 64；寄存器字段','86 B'),('Front-end buffer','8 micro-ops','64 B'),('Reconvergence stack','8 项；每项 6-B PC + 128-bit mask','176 B'),('FLR + LCR + SBB','6 B + 2 B + 1 bit','约 8.1 B'),('Loop-bound detector','两份 16×8-bit 映射快照 + compare/branch 寄存器','48 B'),('VTT','16 bits','2 B'),('NDM IR + ILR','7 bits + 6 B','约 6.9 B'),('合计','论文报告','1139 B')],[2.0,3.65,1.2])
callout('核对','逐项按论文的整数 byte 口径相加为 1139 B。模拟器中的 C++ 对象大小不应被拿来验证该数字；应单独写“理论硬件位数表”。')

heading('7. 基线实验配置',1)
table(['组件','论文配置'],[
('核心','4.0 GHz，OoO，5-wide fetch/dispatch/rename/commit，15 级前端'),('ROB / 队列','ROB 350；issue 128；load 128；store 72'),('分支预测','8 KB TAGE-SC-L（CBP 2016）'),('物理寄存器','256×64-bit integer；256×128-bit FP；128×512-bit vector'),('向量单元','3 ALU，2 shift，2 add，2 mul，2 shuffle'),('L1I','32 KB，4-way，2 cycles'),('L1D','32 KB，8-way，4 cycles，24 MSHRs；16-stream stride prefetcher'),('L2','私有 256 KB，8-way，8 cycles'),('L3','共享 8 MB，16-way，30 cycles'),('内存','最小 50 ns；51.2 GB/s；request-based contention'),('模拟器','Sniper 6.0，最详细 cycle-level core model'),('采样','ROI 跳过初始化；随后 500M representative instructions')],[1.75,5.1])
heading('7.1 Workload 与输入',2)
table(['类别','程序/输入','注意'],[
('GAP','bc, bfs, cc, pr, sssp','每个程序跑 5 张图'),('图输入','Kron, LiveJournal, Orkut, Twitter, Urand','节点/边规模须与论文表 2 对齐'),('HPC/DB','Camel, Graph500, HJ2, HJ8, Kangaroo, NAS-CG, NAS-IS, RandomAccess','版本、编译器、输入规模论文未完整给出')],[1.1,3.15,2.6])
table(['图','节点数（百万）','边数（百万）','论文 LLC MPKI'],[
('Kron','134.2','2111.6','19'),('LiveJournal','4.8','69.0','21'),('Orkut','3.1','1930.3','18'),('Twitter','61.6','1468.4','61'),('Urand','134.2','2147.4','32')],[2.05,1.55,1.55,1.7])

heading('8. Sniper 复现架构建议',1)
callout('工程取舍','Sniper 6.0 年代较早。先固定能运行的 commit、编译器和依赖容器；不要一开始升级到新版并同时改核心模型，否则基线偏差与 DVR bug 会纠缠。')
table(['模块','新增/修改职责','关键接口'],[
('DvrStrideDetector','RPT 更新、候选/inner 检测','dispatch PC；execute effective address'),('DvrDiscovery','VTT/FLR/LCR/SBB、两份映射快照','uop src/dst、branch target/result、rename map'),('DvrController','IDLE→DISCOVERY→SPAWN→RUN/NDM→TERMINATE','主线程 PC、timeout、资源可用性'),('DvrVRAT','scalar/vector 映射、free list、dead-source','rename/execute completion'),('DvrVIR','16 个向量副本的 issued/executed/mask','执行端口、LSQ、vector FU'),('DvrReconvergence','PC/mask 分组与 8 项栈','branch lane results'),('DvrMemory','gather 拆分、prefetch 标记、MSHR/缓存竞争','LSQ/cache hierarchy'),('DvrStats','触发、覆盖、准确率、及时性、MLP、污染','demand/prefetch request lifecycle')],[1.25,3.15,2.7])
heading('8.1 建议状态机',2)
table(['状态','进入条件','退出条件'],[
('IDLE','无 DVR','稳定 stride 候选→DISCOVERY'),('DISCOVERY','候选 stride 触发','再次到候选 PC；无 FLR→IDLE；否则 ARM'),('ARM','已知 chain/bound','主线程再次到 stride PC→RUN 或 NDM'),('RUN','常规 DVR','FLR/next stride/200-uop timeout→CLEANUP'),('NDM','inner remain<64','收集≤128 inner lane→RUN；200-uop fallback'),('CLEANUP','释放临时寄存器与 mask','完成→IDLE')],[1.1,2.9,3.1])
heading('8.2 推荐统计量（必须先加，再调性能）',2)
bullet('dvr_trigger、discovery_success、no_dependent_load、bound_success/fallback、ndm_trigger/fallback、timeout、stack_overflow。')
bullet('generated_lane、active_lane、masked_lane、out_of_bound_lane、vector_uop、scalarized_gather、MSHR_blocked_cycles。')
bullet('prefetch issued / unique / useful / late / evicted-before-use / redundant / wrong-path。')
bullet('主线程 demand latency 分布：L1、L2、L3、off-chip；平均 occupied MSHR（对应论文图 9）。')
bullet('端口争用：DVR 因 main-ready 被拒绝次数；main 因共享 cache/LSQ/MSHR 间接受阻周期。')

heading('9. 分阶段实施计划',1)
table(['阶段','实现内容','通过标准'],[
('P0 基线冻结','Sniper 6.0、编译链、配置、ROI；跑 pointer chase/GAP 小输入','重复 3 次结果一致；基线 IPC/MPKI 存档'),('P1 地址 oracle','离线/理想化预取未来链，验证潜在上限','确实降低 demand miss，形成 Oracle 参考'),('P2 Discovery','RPT、VTT、FLR、loop bound，不发预取','输出的 chain/bound 与 trace 人工核对'),('P3 单层功能 DVR','最多 128 lane，直接事件方式发 prefetch','地址正确、不越界、性能方向正确'),('P4 周期/资源模型','VIR、端口优先、LSQ/MSHR、cache 竞争','关闭资源约束时退化到 P3；开启后无死锁'),('P5 分歧','mask + reconvergence stack','分支 microbench 每条 lane 路径正确'),('P6 NDM','跨 outer invocation flatten lane','短 inner-loop MLP 明显上升'),('P7 论文实验','PRE/IMP/VR/DVR/Oracle；消融、ROB 灵敏度','图 7–12 趋势与关键均值可解释')],[1.15,3.25,2.7])
heading('9.1 功能伪代码',2)
code = '''on_main_load_execute(uop, addr):\n  rpt.update(uop.pc, addr)\n  if state == IDLE and rpt.confident_stride(uop.pc):\n      begin_discovery(uop)\n\non_main_uop(uop):\n  if state != DISCOVERY: return\n  maybe_switch_to_inner_stride(uop)\n  propagate_vtt(uop)\n  if uop.is_load and address_input_tainted(uop): FLR = uop.pc\n  observe_compare_and_backward_branch(uop)\n  if uop.pc == candidate_stride_pc:\n      remaining = infer_bound(checkpoint_in, current_map) or 128\n      state = ARM\n\non_main_reaches_candidate():\n  init_vrat_from_main(); lanes = min(remaining, 128)\n  state = NDM if lanes < 64 else RUN\n\non_dvr_cycle():\n  if main_has_ready_uop_for(port): yield_port()\n  else issue_next_active_vector_copy()\n  if branch_diverges: push_pc_mask_groups()\n  if reached_FlR_or_next_stride() or retired_uops >= 200: terminate()'''
p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(.25); p.paragraph_format.right_indent=Inches(.25); p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(8); border_paragraph(p,GRAY,'F2F4F6')
for i,line in enumerate(code.splitlines()):
    r=p.add_run(line + ('\n' if i<len(code.splitlines())-1 else '')); font(r,8.5,False,DARK,name='Menlo')

heading('10. Microbenchmark 设计',1)
table(['用例','代码形态','预期'],[
('两级间接链','sum += C[B[A[i]]]','VTT 覆盖整条链；FLR=C load；MLP 上升'),('短内循环','CSR graph：每点 2–8 条边','触发 NDM；多 vertex 的 edge 合并至 128 lane'),('精确边界','N=13/63/64/65/127/128/129','active mask 与 min(remaining,128) 一致'),('递减/步长','for(i=N;i>0;i-=2)','bound 与 increment 方向正确'),('分歧','if(A[i]&1) B[...] else C[...]','两组 PC/mask，最终都到终止点'),('重复地址','大量 lane 指向同一 cache line','统计 redundant；不重复占满 MSHR'),('fault/FP 输入','地址链混入不支持值','相关 lane 被 mask，不影响主线程'),('timeout','break/goto 离开识别循环','≤200 DVR 指令后释放全部资源')],[1.25,3.05,2.8])

heading('11. 结果复现与对图方式',1)
table(['论文图','要复现的指标','关键趋势'],[
('图 2','不同 ROB 下 OoO/VR normalized IPC + full-ROB stall%','ROB 越大，VR 机会和收益下降'),('图 7','PRE/IMP/VR/DVR/Oracle normalized IPC','DVR 几何/调和平均约 2.4× baseline；约 2× VR；最高约 6.4×'),('图 8','VR→Offload→+Discovery→+Multiple/NDM 消融','Offload 约从 1.2× 到近 1.5×；完整 DVR 最优'),('图 9','平均每周期占用 MSHR 数','OoO <4；DVR >10（平均量级）'),('图 10','off-chip access 相对量 + normal/runahead 占比','DVR 比 VR 准确且覆盖更广'),('图 11','prefetched line 在 demand 时所在层级','多数在 L1；约 10–20% 仍为 off-chip/未及时'),('图 12','ROB=128/192/224/350/512 的 OoO 与 DVR','DVR 收益保持；论文报告相对各自 OoO 约 1.9/2.2/2.2/2.4/2.5×')],[.75,2.8,3.55])
heading('11.1 归一化与平均',2)
bullet('每个 benchmark-input 的 normalized IPC = IPC_tech / IPC_baseline_same_case。')
bullet('论文最终 H-mean 使用调和平均；复现脚本必须明确是跨所有 33 个 case，还是先对类别平均再做 H-mean。建议两种都输出，图中口径以论文为准。')
bullet('同一 workload 的 ROI 指令数、输入、线程数、编译参数必须完全一致，否则 IPC、MPKI、branch behavior 都会漂移。')
callout('不要只看 speedup','DVR 可能通过大量错误预取“偶然”变快。必须同时报告 accuracy、coverage、timeliness、off-chip traffic、cache pollution 和 MSHR pressure。','warning')

heading('12. 最可能导致复现失败的歧义',1)
table(['歧义','为什么重要','处理方式'],[
('Sniper patch 未公开','流水线与向量执行建模细节不可见','联系作者；所有自行选择写入 assumption log'),('benchmark 版本/命令缺失','输入布局、编译优化改变访存链','保存 commit、编译器、flags、命令行和数据 checksum'),('“主线程无 ready 指令”粒度','可按端口、FU 类型或全局判断','优先按论文原文：same execution port；做替代模型敏感性'),('vector gather 代价','拆分、TLB、cache-bank、带宽模型影响巨大','逐层记录请求；先对齐 24 MSHR，再扩展 TLB/bank'),('寄存器快照语义','论文说 mapping，而非所有值','实现 mapping+读取对应值；记录 checkpoint 时点'),('FLR/复杂 CFG','最后动态 load 不等于静态后支配点','先忠实动态算法；对多路径 case 单独验证'),('物理寄存器压力','论文有 128 vector PR，但 VRAT 同时映射需求复杂','实现阻塞/失败统计，不允许无穷寄存器'),('缓存预取插入策略','prefetch fill、替换优先级可能未写全','默认与普通 load 相同并做替换策略敏感性')],[1.45,2.35,3.3])

heading('13. 验收清单',1)
heading('正确性',2)
bullet('□ DVR 不修改主线程体系结构寄存器、内存或控制流。')
bullet('□ 每个预取地址可追溯到 stride seed、lane、taint chain 和动态 PC。')
bullet('□ bound 成功时不生成超出剩余迭代数的 active lane。')
bullet('□ fault、无效 source、FP 不支持、分歧未选路径均被 mask。')
bullet('□ 超时、栈满、寄存器不足、MSHR 满时可安全退让/终止，无死锁。')
heading('实验完整性',2)
bullet('□ 基线 stride prefetcher 始终开启；DVR 不是替代它。')
bullet('□ 所有方案共享完全相同的 core/cache/memory/ROI 配置。')
bullet('□ 至少包含 VR、DVR、Oracle；理想情况下还包含 PRE 与 IMP。')
bullet('□ 运行消融：VR、Offload、+Discovery、+NDM。')
bullet('□ 报告 ROB 128/192/224/350/512 敏感性，并按比例扩展 backend structures。')
bullet('□ 结果脚本、raw stats、配置、日志、seed、commit 都归档。')

heading('14. 建议目录结构与实验记录',1)
tree='''dvr-repro/\n  simulator/              # 固定 Sniper 6.0 与 DVR patch\n  configs/                # baseline / vr / dvr / oracle / ROB sweep\n  benchmarks/             # build scripts；不提交受限数据\n  inputs/manifest.yaml    # URL、版本、checksum、预处理命令\n  microbench/             # indirect、nested、divergence、bound tests\n  scripts/run.py          # 生成任务与 ROI 命令\n  scripts/collect.py      # IPC、MLP、traffic、accuracy、timeliness\n  results/raw/            # 原始模拟输出\n  results/tidy/           # 一行一个 case 的 CSV/Parquet\n  plots/                  # 图 2、7–12\n  assumptions.md          # 论文未明确处的决策\n  environment.lock        # compiler、依赖、容器摘要'''
p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(.25); border_paragraph(p,GRAY,'F2F4F6')
for i,line in enumerate(tree.splitlines()): font(p.add_run(line+('\n' if i<len(tree.splitlines())-1 else '')),8.4,False,DARK,name='Menlo')
heading('实验记录最少字段',2)
para('run_id, date, host, git_commit, simulator_version, compiler, flags, benchmark, input, ROI, configuration, seed, wall_time, simulated_instructions, cycles, IPC, L1/L2/L3 MPKI, branch MPKI, average_MSHR, offchip_requests, DVR counters, exit_status。')

heading('15. 论文结果的正确解读',1)
para('DVR 的 2.4× 并不意味着现实 CPU 加 1139 B 就能直接得到 2.4×。1139 B 是论文枚举的控制结构位数，不包含共享向量执行单元、寄存器文件、LSQ/MSHR、cache/bandwidth 的既有成本，也不等同于物理设计后的面积/功耗。实验是周期级模拟，并对一组高度内存受限、间接访问密集的 workload 取平均。')
para('真正可迁移的结论是：对于可由规则 seed load 引出的不规则依赖链，“主动、解耦、跨迭代向量化”比“等大 ROB 堵住再 runahead”更适合持续制造 MLP；而准确性依赖 loop-bound detection，短内循环则需要跨 invocation 聚合。')
callout('研究延伸','如果严格复现受阻，可把工作转成一篇可发表的复现/再评估：在 gem5 或更新 Sniper 上实现 DVR，量化 TLB、cache pollution、带宽饱和、能耗、不同 vector width 与更现代 ROB/内存系统下的收益。','note')

heading('附录 A｜作者联系清单',1)
bullet('请求 MICRO ’23 使用的 Sniper commit 与 DVR patch。')
bullet('请求 PRE、IMP、VR、Oracle 的实现版本和配置差异。')
bullet('请求 13 个 benchmark 的源码 commit、编译器/flags、线程数、输入生成脚本和 ROI 标记位置。')
bullet('请求图 7–12 的原始 CSV、平均口径以及 ROB sweep 时 backend structures 的精确缩放规则。')
bullet('询问 gather 拆分、cache fill/replacement、TLB、MSHR 冲突与端口仲裁的未写细节。')

heading('附录 B｜来源与引用',1)
para('主论文：A. Naithani, J. Roelandts, S. Ainsworth, T. M. Jones, and L. Eeckhout, “Decoupled Vector Runahead,” MICRO ’23, pp. 17–31, DOI: 10.1145/3613424.3614255。本文所有机制、数字与论文图号均依据用户提供的 PDF。')
para('公开条目：Ghent University Academic Bibliography, “Decoupled vector runahead,” http://hdl.handle.net/1854/LU-01HMR8HPGNZS1RC118QA1GPBMG。')
para('作者公开页面与 MICRO 2023 program 用于核对论文身份；未发现该论文的公开 artifact 链接。检索结论应视为“公开检索未找到”，不等于代码一定不存在。')

doc.core_properties.title='Decoupled Vector Runahead：论文解读与复现指南'
doc.core_properties.subject='MICRO 2023 DVR 机制与 Sniper 复现路线'
doc.core_properties.author='Codex'
doc.core_properties.keywords='DVR, Vector Runahead, Sniper, Prefetching, MLP'
doc.save(OUT)
print(OUT)
